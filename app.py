"""
Facebook Messenger Bot សម្រាប់ឆ្លើយសំណួរសិស្ស BBU
=====================================================
រចនាសម្ព័ន្ធ៖
1. Flask web server ទទួល webhook ពី Facebook
2. អាន Knowledge Base ពី knowledge/*.txt, *.docx, *.pdf ដាក់ចូល Prompt ដោយផ្ទាល់
3. ផ្ញើទៅ Google Gemini API ដើម្បីតែងចម្លើយជាភាសាធម្មជាតិ
4. ផ្ញើចម្លើយត្រឡប់ទៅសិស្សតាម Facebook Send API
"""

import os
import hmac
import hashlib
import glob
import re
import requests
from flask import Flask, request, jsonify
from google import genai
from google.genai import types
from docx import Document
from pypdf import PdfReader

# ---------- ១. ការកំណត់រចនាសម្ព័ន្ធ (Config) ----------
# តម្លៃទាំងនេះមកពី Environment Variables (កុំដាក់ត្រង់នេះដោយផ្ទាល់ ពេល deploy ពិត)
VERIFY_TOKEN = os.environ.get("VERIFY_TOKEN", "readclub_bbu_verify_2026")
PAGE_ACCESS_TOKEN = os.environ.get("PAGE_ACCESS_TOKEN", "")
APP_SECRET = os.environ.get("APP_SECRET", "")
GEMINI_API_KEY = os.environ.get("GEMINI_API_KEY", "")

app = Flask(__name__)
gemini_client = genai.Client(api_key=GEMINI_API_KEY)
GEMINI_MODEL_NAME = "gemini-3.6-flash"

# ---------- ២. ផ្ទុក Knowledge Base ចូល memory ពេល server ចាប់ផ្តើម ----------
def read_docx_text(filepath):
    """អាកអត្ថបទទាំងអស់ចេញពីឯកសារ Word (.docx) — កថាខណ្ឌ + តារាង"""
    doc = Document(filepath)
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip(" |"):
                parts.append(row_text)
    return "\n".join(parts)


def read_pdf_text(filepath):
    """អានអត្ថបទចេញពីឯកសារ PDF (ទំព័រនីមួយៗ)។ PDF ជា Scan/រូបភាព នឹងអានមិនចេញអត្ថបទទេ"""
    reader = PdfReader(filepath)
    parts = []
    for page in reader.pages:
        page_text = page.extract_text() or ""
        if page_text.strip():
            parts.append(page_text.strip())
    return "\n".join(parts)


def extract_qa_blocks(content):
    """ញែកទម្រង់ 'Q: ... \\n A: ...' ចេញពីអត្ថបទ"""
    blocks = re.findall(r"Q:\s*(.+?)\s*\nA:\s*(.+?)(?=\n\n|\nQ:|\Z)", content, re.DOTALL)
    qa_texts = []
    for question, answer in blocks:
        answer = answer.strip()
        if answer.startswith("[") and answer.endswith("]"):
            continue  # រំលងចម្លើយគំរូដែលមិនទាន់បំពេញ
        qa_texts.append(f"សំណួរ: {question.strip()}\nចម្លើយ: {answer}")
    return qa_texts


def load_knowledge_base(folder="knowledge"):
    """
    អានឯកសារ .txt, .docx, .pdf ទាំងអស់ក្នុង folder knowledge/
    - បើឯកសារមានទម្រង់ 'Q: ... A: ...' នឹងញែកជា Q&A ដាច់ដោយឡែក
    - បើគ្មានទម្រង់នេះ (ឧ. ឯកសារ Word/PDF ជាអត្ថបទសេរី) នឹងយកអត្ថបទទាំងមូលធ្វើជាព័ត៌មានយោង
    """
    text_blocks = []

    txt_files = glob.glob(os.path.join(folder, "*.txt"))
    docx_files = glob.glob(os.path.join(folder, "*.docx"))
    pdf_files = glob.glob(os.path.join(folder, "*.pdf"))

    for filepath in txt_files:
        with open(filepath, "r", encoding="utf-8") as f:
            content = f.read()
        qa_texts = extract_qa_blocks(content)
        if qa_texts:
            text_blocks.extend(qa_texts)
        else:
            text_blocks.append(content.strip())

    for filepath in docx_files:
        # រំលងឯកសារបណ្តោះអាសន្នដែល Word បង្កើតស្វ័យប្រវត្តិ (ឈ្មោះចាប់ផ្តើមដោយ ~$)
        if os.path.basename(filepath).startswith("~$"):
            continue
        try:
            content = read_docx_text(filepath)
        except Exception as e:
            print(f"មិនអាចអានឯកសារ {filepath}:", e)
            continue
        qa_texts = extract_qa_blocks(content)
        if qa_texts:
            text_blocks.extend(qa_texts)
        else:
            text_blocks.append(content.strip())

    for filepath in pdf_files:
        try:
            content = read_pdf_text(filepath)
        except Exception as e:
            print(f"មិនអាចអានឯកសារ {filepath}:", e)
            continue
        if not content.strip():
            print(f"ព្រមាន៖ {filepath} មិនមានអត្ថបទអានចេញបានទេ (ប្រហែលជា PDF ជា Scan/រូបភាព)")
            continue
        qa_texts = extract_qa_blocks(content)
        if qa_texts:
            text_blocks.extend(qa_texts)
        else:
            text_blocks.append(content.strip())

    all_files = txt_files + [f for f in docx_files if not os.path.basename(f).startswith("~$")] + pdf_files
    print(f"[Knowledge Base] ឯកសារដែលរកឃើញ ({len(all_files)}): {[os.path.basename(f) for f in all_files]}")
    print(f"[Knowledge Base] ចំនួន block សរុប: {len(text_blocks)}")

    return "\n\n".join(t for t in text_blocks if t)


knowledge_base_text = load_knowledge_base()
print(f"[Knowledge Base] ចំនួនអក្សរសរុបដែលផ្ទុកបាន: {len(knowledge_base_text)}")
print("[Knowledge Base] ១៥០០ តួអក្សរដំបូង:\n" + knowledge_base_text[:1500])


# ---------- ៣. ហៅ Gemini API ដើម្បីតែងចម្លើយ ----------
def call_gemini(prompt_text):
    """ហៅ Gemini ជាមួយ thinking_level='low' ដើម្បីកុំឲ្យ 'ការគិត'ស៊ីអស់ Token budget"""
    return gemini_client.models.generate_content(
        model=GEMINI_MODEL_NAME,
        contents=prompt_text,
        config=types.GenerateContentConfig(
            max_output_tokens=800,
            thinking_config=types.ThinkingConfig(thinking_level="low"),
        ),
    )


def generate_answer(user_message, debug=False):
    if knowledge_base_text:
        system_prompt = (
            "អ្នកគឺជាជំនួយការឆ្លើយសំណួរសម្រាប់សិស្សថ្នាក់ទី១២ដែលចាប់អារម្មណ៍ចូលរៀននៅ "
            "Build Bright University (BBU)។ ប្រើព័ត៌មានខាងក្រោមដើម្បីឆ្លើយសំណួរអ្នកប្រើប្រាស់ "
            "ជាភាសាខ្មែរ ដោយសង្ខេប ច្បាស់លាស់ និងគួរសម។ "
            "សរសេរចម្លើយឡើងវិញជាពាក្យផ្ទាល់ខ្លួន កុំចម្លងឃ្លាវែងៗពីព័ត៌មានយោងដោយផ្ទាល់ "
            "(សង្ខេប/សរសេរសារជាថ្មីជំនួសការចម្លង)។ "
            "បើព័ត៌មានមិនគ្រប់គ្រាន់ សូមណែនាំឲ្យទាក់ទងការិយាល័យចុះឈ្មោះដោយផ្ទាល់។\n\n"
            f"ព័ត៌មានយោង៖\n{knowledge_base_text}"
        )
    else:
        system_prompt = (
            "អ្នកគឺជាជំនួយការឆ្លើយសំណួរសម្រាប់សិស្សថ្នាក់ទី១២ដែលចាប់អារម្មណ៍ចូលរៀននៅ "
            "Build Bright University (BBU)។ ចម្លើយជាភាសាខ្មែរ ដោយសង្ខេប និងគួរសម។ បើមិនប្រាកដច្បាស់ "
            "សូមណែនាំឲ្យទាក់ទងការិយាល័យចុះឈ្មោះដោយផ្ទាល់ ជំនួសការស្មានចម្លើយ។"
        )

    full_prompt = f"{system_prompt}\n\nសំណួរអ្នកប្រើប្រាស់៖ {user_message}"
    response = call_gemini(full_prompt)

    # ពិនិត្យមូលហេតុបញ្ឈប់ (finish_reason) — ជួយវិនិច្ឆ័យពេលចម្លើយដាច់កណ្តាល
    finish_reason = None
    try:
        finish_reason = response.candidates[0].finish_reason.name
    except Exception:
        pass

    if debug:
        return {"text": getattr(response, "text", "") or "", "finish_reason": finish_reason}

    if finish_reason == "RECITATION":
        # Gemini បញ្ឈប់ព្រោះកំពុងចម្លងអត្ថបទដើមច្រើនពេក — សាកល្បងម្តងទៀតដោយសុំសង្ខេបខ្លីជាង
        retry_prompt = full_prompt + "\n\n(សូមឆ្លើយខ្លីជាងនេះ ដោយសរសេរជាពាក្យផ្ទាល់ខ្លួន កុំចម្លងឃ្លាវែងៗពីព័ត៌មានយោង)"
        response = call_gemini(retry_prompt)

    try:
        return response.text
    except Exception:
        return "សូមទោស! ខ្ញុំមិនអាចបង្កើតចម្លើយពេញលេញបានទេ សូមសាកសួរម្តងទៀតដោយប្រើសំណួរខ្លីជាង ឬទាក់ទងការិយាល័យដោយផ្ទាល់។"


# ---------- ៤. ផ្ញើសារត្រឡប់ទៅ Facebook ----------
def send_message(recipient_id, message_text):
    url = "https://graph.facebook.com/v21.0/me/messages"
    params = {"access_token": PAGE_ACCESS_TOKEN}
    payload = {
        "recipient": {"id": recipient_id},
        "message": {"text": message_text},
    }
    resp = requests.post(url, params=params, json=payload)
    if resp.status_code != 200:
        print("Send API error:", resp.text)


# ---------- ៥. Webhook Verification (GET request ពី Facebook) ----------
@app.route("/webhook", methods=["GET"])
def verify_webhook():
    mode = request.args.get("hub.mode")
    token = request.args.get("hub.verify_token")
    challenge = request.args.get("hub.challenge")

    if mode == "subscribe" and token == VERIFY_TOKEN:
        return challenge, 200
    return "Verification failed", 403


# ---------- ៦. ការត្រួតពិនិត្យសុវត្ថិភាព Signature ----------
def verify_signature(payload, signature_header):
    """ត្រួតពិនិត្យថា request នេះមកពី Facebook ពិតប្រាកដ (ការពារ spoofing)"""
    if not APP_SECRET or not signature_header:
        return False
    expected_hash = hmac.new(
        APP_SECRET.encode("utf-8"), payload, hashlib.sha256
    ).hexdigest()
    received_hash = signature_header.replace("sha256=", "")
    return hmac.compare_digest(expected_hash, received_hash)


# ---------- ៧. ទទួលសារពីសិស្ស (POST request ពី Facebook) ----------
@app.route("/webhook", methods=["POST"])
def handle_message():
    signature = request.headers.get("X-Hub-Signature-256", "")
    if not verify_signature(request.get_data(), signature):
        return jsonify({"status": "invalid signature"}), 403

    data = request.get_json()

    if data.get("object") == "page":
        for entry in data.get("entry", []):
            for messaging_event in entry.get("messaging", []):
                sender_id = messaging_event["sender"]["id"]

                if "message" in messaging_event and "text" in messaging_event["message"]:
                    user_text = messaging_event["message"]["text"]
                    try:
                        answer = generate_answer(user_text)
                    except Exception as e:
                        print("Error generating answer:", e)
                        answer = "សូមទោស! ឥឡូវនេះមានបញ្ហាបច្ចេកទេស សូមព្យាយាមម្តងទៀត ឬទាក់ទងផ្ទាល់។"
                    send_message(sender_id, answer)

    return jsonify({"status": "ok"}), 200


# ---------- ៨. Health check (សម្រាប់ Render ដឹងថា server រស់) ----------
@app.route("/", methods=["GET"])
def health_check():
    return "BBU Facebook Bot is running.", 200


# ---------- ៩. Debug — មើល Knowledge Base ដែលផ្ទុកបានពិត (សម្រាប់ត្រួតពិនិត្យប៉ុណ្ណោះ) ----------
@app.route("/debug/knowledge", methods=["GET"])
def debug_knowledge():
    return {
        "total_characters": len(knowledge_base_text),
        "content": knowledge_base_text,
    }, 200


# ---------- ១០. Debug — សាកល្បងសួរសំណួរដោយផ្ទាល់តាម Browser (មិនចាំបាច់ឆ្លងកាត់ Facebook) ----------
@app.route("/debug/ask", methods=["GET"])
def debug_ask():
    question = request.args.get("q", "")
    if not question:
        return "សូមដាក់សំណួរជាមួយ ?q=សំណួររបស់អ្នក ក្នុង URL", 400
    try:
        result = generate_answer(question, debug=True)
    except Exception as e:
        return f"Error: {e}", 500
    return {"question": question, "answer": result["text"], "finish_reason": result["finish_reason"]}, 200


if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port)
